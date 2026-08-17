"""
Generate a comprehensive project documentation Word document for
vedic-astrology-react.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os

# ── helpers ──────────────────────────────────────────────────────────────────


def set_cell_shading(cell, hex_color: str):
    """Set background shading on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Set borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "D4A574")
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def heading1(doc, text):
    p = doc.add_heading(text, level=1)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x7B, 0x3F, 0x00)
    run.font.size = Pt(22)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    return p


def heading2(doc, text):
    p = doc.add_heading(text, level=2)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x9B, 0x5B, 0x00)
    run.font.size = Pt(16)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p


def heading3(doc, text):
    p = doc.add_heading(text, level=3)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)
    run.font.size = Pt(13)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(doc, text, bold=False, italic=False):
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_after = Pt(3)
    return p


def make_table(doc, headers, rows, col_widths=None):
    """Create a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "7B3F00")

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            if ri % 2 == 1:
                set_cell_shading(cell, "FDF5E6")

    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[ci])

    return table


def add_divider(doc):
    """Add a horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D4A574")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── document setup ────────────────────────────────────────────────────────────

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.15)
section.right_margin = Inches(1.15)

# Title page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
run = p.add_run("\n\n\n")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Vedic Astrology")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x7B, 0x3F, 0x00)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Project Documentation")
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x9B, 0x5B, 0x00)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run("Comprehensive Overview — Features, Architecture & Deliverables")
run.font.size = Pt(14)
run.font.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run("Built for: Sampath Kumara — Vedic Astrologer")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
run = p.add_run(
    "Technologies: React 19 + TypeScript + Tailwind CSS + Express + Supabase"
)
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
run = p.add_run("Location: Rajarajeshwari Nagar, Bengaluru — 560098")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ── TABLE OF CONTENTS ─────────────────────────────────────────────────────────
heading1(doc, "Table of Contents")
toc_items = [
    ("1", "Project Overview", ""),
    ("2", "Technology Stack", ""),
    ("3", "Project Architecture", ""),
    ("4", "Public-Facing Pages & Features", ""),
    ("  4.1", "Homepage", ""),
    ("  4.2", "About Us", ""),
    ("  4.3", "Astrology Consultations Services", ""),
    ("  4.4", "Homams — Sacred Vedic Fire Rituals", ""),
    ("  4.5", "Birth Chart PDF Report", ""),
    ("  4.6", "Chat with Guruji (AI Assistant)", ""),
    ("  4.7", "Instant Palm Reading", ""),
    ("  4.8", "Testimonials", ""),
    ("  4.9", "Contact Us", ""),
    ("  4.10", "Legal Pages", ""),
    ("5", "Admin CMS (Content Management System)", ""),
    ("6", "Backend API Services", ""),
    ("7", "Database & Authentication (Supabase)", ""),
    ("8", "Key Integrations", ""),
    ("9", "SEO & Performance Features", ""),
    ("10", "Design & UX Features", ""),
    ("11", "Deployment & Build Configuration", ""),
    ("12", "Environment Variables & Configuration", ""),
    ("13", "Testing & Quality Assurance", ""),
    ("14", "Summary of Deliverables", ""),
]
tbl = doc.add_table(rows=len(toc_items), cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (num, title, sub) in enumerate(toc_items):
    row = tbl.rows[i]
    row.cells[0].text = ""
    row.cells[1].text = ""
    row.cells[2].text = ""
    p0 = row.cells[0].paragraphs[0]
    run0 = p0.add_run(num)
    run0.font.size = Pt(11)
    p1 = row.cells[1].paragraphs[0]
    run1 = p1.add_run(title)
    run1.font.size = Pt(11)
    run1.font.bold = sub == ""
    run1.font.color.rgb = (
        RGBColor(0x7B, 0x3F, 0x00) if sub == "" else RGBColor(0x33, 0x33, 0x33)
    )
    p2 = row.cells[2].paragraphs[0]
    run2 = p2.add_run(sub)
    run2.font.size = Pt(10)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "1. Project Overview")
add_divider(doc)

body(
    doc,
    (
        "The Vedic Astrology web application is a complete, production-ready digital platform "
        "built for an experienced Vedic astrologer — Sampath Kumara — based in Bengaluru, India. "
        "The website serves as a fully functional online presence that enables visitors from around "
        "the world to access authentic Vedic astrology services, book sacred rituals, receive "
        "personalized astrological guidance, and engage directly with the astrologer."
    ),
)

heading2(doc, "1.1 Purpose & Vision")
body(
    doc,
    (
        "The platform bridges ancient Vedic wisdom with modern technology, making authentic "
        "Vedic astrological guidance accessible to anyone with an internet connection. The vision "
        "is to provide honest, grounded spiritual guidance — not fortune-telling — rooted in "
        "traditional Vedic principles including natal chart (kundli) analysis, planetary dosha "
        "interpretation, muhurta (auspicious timing), sacred homam rituals, and personalized "
        "remedies."
    ),
)

heading2(doc, "1.2 Key Objectives")
bullet(
    doc,
    "Deliver authentic Vedic astrology consultations for a wide range of life concerns",
)
bullet(
    doc, "Enable online booking of 15+ sacred Vedic homam rituals with payment support"
)
bullet(
    doc,
    "Provide AI-powered first-contact guidance via a chatbot trained in astrological scope",
)
bullet(doc, "Generate detailed Vedic birth chart PDF reports with dosha analysis")
bullet(doc, "Offer instant palm reading using AI image analysis")
bullet(
    doc,
    "Build trust through client testimonials, detailed service descriptions, and legal transparency",
)
bullet(
    doc,
    "Provide a secure admin CMS for managing all content without developer assistance",
)
bullet(
    doc,
    "Ensure legal compliance with Privacy Policy, Terms & Conditions, Disclaimer, and Refund policy",
)
bullet(doc, "Optimize for search engines (SEO) and social media sharing")

heading2(doc, "1.3 Brand Details")
make_table(
    doc,
    ["Attribute", "Value"],
    [
        ["Brand Name", "Vedic Astrology"],
        ["Astrologer", "Sampath Kumara"],
        [
            "Location",
            "No 100, 21st Main, 8th Cross, Rajarajeshwari Nagar, Bengaluru — 560098",
        ],
        ["Phone", "+91 98861 00565"],
        ["WhatsApp", "+91 98861 00565"],
        ["Email", "guruji@vedicastrology.com"],
        ["Working Hours", "Mon – Sun, 8:00 AM – 9:00 PM IST"],
        ["Tagline", "Ancient Wisdom • Cosmic Guidance"],
    ],
    col_widths=[2.5, 4.0],
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "2. Technology Stack")
add_divider(doc)

heading2(doc, "2.1 Frontend Technologies")
make_table(
    doc,
    ["Technology", "Version", "Purpose"],
    [
        ["React", "19.2", "UI component library for building the user interface"],
        ["TypeScript", "~5.8", "Type-safe JavaScript for robust development"],
        ["React Router", "v6.30", "Client-side routing for SPA navigation"],
        ["Tailwind CSS", "v4", "Utility-first CSS framework for rapid styling"],
        ["PostCSS", "v8", "CSS processing pipeline for Tailwind v4"],
        ["Vite", "v6.3", "Fast build tool and development server"],
        [
            "Framer Motion",
            "v12.43",
            "Animation library for smooth page transitions and effects",
        ],
        ["React Hook Form", "v7.83", "Form state management and validation"],
        ["Zod", "v3.24", "Runtime schema validation for forms"],
        ["Lucide React", "v1.23", "Icon library"],
        ["React Helmet Async", "v3.0", "Dynamic document head management for SEO"],
    ],
    col_widths=[2.0, 1.0, 3.5],
)

heading2(doc, "2.2 Backend Technologies")
make_table(
    doc,
    ["Technology", "Version", "Purpose"],
    [
        ["Express.js", "v5.1", "REST API server (port 3001)"],
        ["TSX", "v4.19", "TypeScript execution for server code with hot-reload"],
        [
            "Google Gemini AI",
            "v0.24",
            "AI engine for chatbot and palm reading features",
        ],
        ["Cookie", "v1.0", "Server-side cookie handling"],
        ["Cookie Parser", "v1.4", "Cookie parsing middleware"],
        ["CORS", "v2.8", "Cross-origin resource sharing configuration"],
        ["UUID", "v11.1", "Unique identifier generation"],
    ],
    col_widths=[2.0, 1.0, 3.5],
)

heading2(doc, "2.3 Database & Auth")
make_table(
    doc,
    ["Technology", "Purpose"],
    [
        [
            "Supabase (PostgreSQL)",
            "Primary database for storing services, homams, pages, and site content",
        ],
        ["Supabase Auth", "Authentication for admin CMS access"],
        ["Supabase SSR", "Server-side rendering support for Supabase clients"],
        ["Supabase JS Client", "Client-side database queries and real-time features"],
    ],
    col_widths=[3.0, 3.5],
)

heading2(doc, "2.4 Developer Tools")
make_table(
    doc,
    ["Tool", "Purpose"],
    [
        ["TypeScript", "Static type checking across frontend and server"],
        ["ESLint", "Code linting and style enforcement"],
        ["Vite Plugin React", "React fast refresh during development"],
    ],
    col_widths=[3.0, 3.5],
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — PROJECT ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "3. Project Architecture")
add_divider(doc)

body(
    doc,
    (
        "The project follows a clean client-server architecture with a React SPA frontend "
        "served by Vite and an Express backend API. Supabase provides the database and "
        "authentication layer."
    ),
)

heading2(doc, "3.1 Architecture Diagram (Conceptual)")
body(
    doc,
    (
        "┌─────────────────────────────────────────────────────────────┐\n"
        "│                        BROWSER / CLIENT                     │\n"
        "│  ┌─────────────┐    ┌──────────────┐    ┌────────────────┐ │\n"
        "│  │  React SPA  │    │  Vite Dev    │    │  Express API   │ │\n"
        "│  │  (Port 5173)│    │  Server       │    │  (Port 3001)   │ │\n"
        "│  └──────┬──────┘    └──────────────┘    └───────┬────────┘ │\n"
        "│         │          Proxy /api → :3001           │           │\n"
        "└─────────┼────────────────────────────────────────┼───────────┘\n"
        "          │                                       │\n"
        "          ▼                                       ▼\n"
        "  ┌──────────────────────────────────────────────────────┐\n"
        "  │              Supabase (PostgreSQL + Auth)            │\n"
        "  │  • Services data    • Homams data                    │\n"
        "  │  • Page content     • Admin users                    │\n"
        "  └──────────────────────────────────────────────────────┘\n\n"
        "  ┌──────────────────────────────────────────────────────┐\n"
        "  │              Google Gemini API                       │\n"
        "  │  • AI Chatbot     • Palm Reading Analysis            │\n"
        "  └──────────────────────────────────────────────────────┘"
    ),
)

heading2(doc, "3.2 Directory Structure")
body(doc, "The project root contains the following key directories and files:")

make_table(
    doc,
    ["Path", "Description"],
    [
        ["src/", "Frontend source code — React components, pages, hooks, styles"],
        ["src/pages/", "14 public-facing page components + 1 Not-Found page"],
        [
            "src/admin/",
            "8 admin CMS pages (login, dashboard, CRUD for services/homams/pages)",
        ],
        [
            "src/components/",
            "Reusable UI components (layout, AI widgets, cards, effects, forms, SEO)",
        ],
        [
            "src/lib/",
            "Core logic — data files, API client, Supabase setup, site config, utilities",
        ],
        [
            "src/lib/data/",
            "Static content — services.ts, homams.ts, faqs.ts, testimonials.ts, legal.ts",
        ],
        ["src/lib/site/", "Site configuration (brand, nav, footer links)"],
        ["server/", "Express backend — API routes, middleware, server utilities"],
        ["server/routes/", "API endpoints — chat, palm-reading, enquiry, admin CRUD"],
        ["server/middleware/", "Admin authentication middleware"],
        [
            "server/lib/",
            "Server-side helpers — Gemini config, rate limiting, chat policy",
        ],
        ["supabase/", "Supabase project config, edge functions, migrations"],
        ["public/", "Static assets — images, icons, PWA service worker"],
        ["scripts/", "Dev orchestrator, build scripts, sitemap generator"],
        ["index.html", "Vite SPA HTML shell"],
        ["package.json", "Dependencies, scripts, engine versions"],
        ["tsconfig.json", "TypeScript configuration with path aliases"],
        ["vite.config.ts", "Vite build config with API proxy and path aliases"],
        ["postcss.config.js", "Tailwind CSS v4 PostCSS plugin config"],
    ],
    col_widths=[2.5, 4.0],
)

heading2(doc, "3.3 Request Flow")
body(doc, "1. User visits the site → Vite serves the React SPA")
body(doc, "2. React Router handles all navigation without full page reloads")
body(doc, "3. API calls to /api/* routes are proxied by Vite to Express on port 3001")
body(
    doc, "4. Express routes handle AI chat, palm reading, contact form, and admin CRUD"
)
body(
    doc,
    "5. Supabase is used for persistent data (services, homams, page content, admin auth)",
)
body(doc, "6. Google Gemini API is called server-side for AI-powered features")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — PUBLIC-FACING PAGES & FEATURES
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "4. Public-Facing Pages & Features")
add_divider(doc)

body(
    doc,
    (
        "The application provides 14 public-facing pages accessible via React Router. "
        "All pages are styled with Tailwind CSS, animated with Framer Motion, and optimized for SEO."
    ),
)

heading2(doc, "4.1 Homepage (/)")
body(
    doc,
    "The homepage is the primary landing page designed to create a strong first impression and guide visitors to key services.",
)
bullet(doc, "Hero section with brand introduction and call-to-action buttons")
bullet(doc, "Service overview cards linking to detailed consultation pages")
bullet(doc, "Featured homams section with visual cards and booking CTAs")
bullet(doc, "Interactive AI chatbot widget (Chat with Guruji) embedded directly")
bullet(doc, "Client testimonials carousel with filtering and star ratings")
bullet(doc, "FAQ accordion covering booking, confidentiality, and general questions")
bullet(doc, "Contact information, map location, and social media links")
bullet(doc, "Scroll-triggered animations for enhanced visual engagement")
bullet(doc, "Floating WhatsApp button for instant messaging")
bullet(doc, "SEO-optimized with dynamic meta tags via React Helmet Async")

heading2(doc, "4.2 About Us (/about-us)")
bullet(doc, "Astrologer biography and credentials for Sampath Kumara")
bullet(doc, "Philosophy and approach to Vedic astrology")
bullet(doc, "Years of experience and specializations")
bullet(doc, "Testimonials showcase")
bullet(doc, "Location details with embedded map")
bullet(doc, "Contact CTA and social media links")

heading2(doc, "4.3 Astrology Consultations Services (/services)")
body(
    doc,
    "The platform offers 10 detailed astrology consultation services, each with its own dedicated page:",
)

make_table(
    doc,
    ["#", "Service", "Price (INR)", "Duration"],
    [
        ["1", "Love & Relationship Problems", "501", "20–30 min"],
        ["2", "Marriage Delay & Divorce Issues", "751", "25–35 min"],
        ["3", "Career Confusion & Job Problems", "751", "25–35 min"],
        ["4", "Financial Instability & Debt Problems", "999", "30 min"],
        ["5", "Family Conflicts & Domestic Issues", "501", "20–30 min"],
        ["6", "Mental Stress, Anxiety & Depression", "501", "25 min"],
        ["7", "Health & Wellness Astrology", "751", "25–30 min"],
        ["8", "Education & Exam Success", "501", "20–25 min"],
        ["9", "Business Growth & Partnership Problems", "1,100", "30–40 min"],
        ["10", "Property, Legal & Court Case Guidance", "1,100", "30–40 min"],
    ],
    col_widths=[0.5, 3.0, 1.2, 1.0],
)

body(
    doc,
    (
        "Each service page includes: detailed description, problem statement, planetary analysis breakdown, "
        "remedies list, benefits, and FAQ section with booking instructions."
    ),
)
bullet(doc, "Dynamic routing: /services/:slug for each service")
bullet(doc, "Enquiry form integrated on every service page")
bullet(doc, "AI chat widget context-aware to the current service topic")
bullet(doc, "WhatsApp quick-contact button per service")

heading2(doc, "4.4 Homams — Sacred Vedic Fire Rituals (/homams)")
body(
    doc,
    "The platform lists 15 authentic Vedic homam (fire ritual) offerings, each with its own dedicated page:",
)

make_table(
    doc,
    ["#", "Homam", "Duration"],
    [
        ["1", "Ganapathi Homam", "1.5 – 2 hours"],
        ["2", "Navagraha Homam", "2 – 3 hours"],
        ["3", "Lakshmi Kubera Homam", "2 – 3 hours"],
        ["4", "Surya Homam", "1.5 – 2 hours"],
        ["5", "Chandra Homam", "1.5 – 2 hours"],
        ["6", "Rudra Homam", "2 – 3 hours"],
        ["7", "Maha Mrityunjaya Homam", "2 – 3 hours"],
        ["8", "Saraswati Homam", "1.5 – 2 hours"],
        ["9", "Durga Homam", "2 – 3 hours"],
        ["10", "Sudarshana Homam", "2 – 3 hours"],
        ["11", "Dhanvantari Homam", "2 – 3 hours"],
        ["12", "Ayushya Homam", "2 hours"],
        ["13", "Shani Shanti Homam", "2 hours"],
        ["14", "Rahu Ketu Shanti Homam", "2 – 3 hours"],
        ["15", "Mangal Dosha Homam", "2 hours"],
        ["16", "Kadali Vivaha", "2 – 3 hours"],
        ["17", "Kumbha Vivaha", "2 – 3 hours"],
        ["18", "Moksha Narayana Bali (Tila Homa)", "3 – 4 hours"],
    ],
    col_widths=[0.5, 3.0, 2.0],
)

body(doc, "Each homam detail page includes:")
bullet(doc, "Full ritual description and spiritual significance")
bullet(doc, "List of benefits and blessings")
bullet(doc, "Suitable for / recommended situations")
bullet(doc, "Pooja items included")
bullet(
    doc, "Booking instructions and payment options (Razorpay, UPI, screenshot upload)"
)
bullet(doc, "FAQ section addressing common questions")

heading2(doc, "4.5 Birth Chart PDF Report (/birth-chart-pdf)")
body(
    doc,
    (
        "A dedicated page where users can request a comprehensive Vedic birth chart (kundli) PDF report. "
        "Users submit their birth details (name, date, exact time, place of birth) via an enquiry form. "
        "The report includes:"
    ),
)
bullet(doc, "Lagna chart (ascendant chart)")
bullet(doc, "Moon sign and Sun sign")
bullet(doc, "Nakshatra (lunar mansion) at birth")
bullet(doc, "Complete planetary positions")
bullet(doc, "Dosha overview (Manglik, Kaal Sarp, etc.)")
bullet(doc, "Indications for career, marriage, finance, and health")
bullet(doc, "Suggested remedies and general life guidance")
body(doc, "Report is delivered via email or WhatsApp within 24–48 hours.")

heading2(doc, "4.6 Chat with Guruji — AI Assistant (/chat-with-guruji)")
body(
    doc,
    (
        "A full-page AI-powered chat experience built on Google's Gemini API. "
        "This is a key differentiator — offering instant, free initial guidance to visitors."
    ),
)
bullet(
    doc,
    "Free-tier model: 3 free messages per visitor (tracked via cookie + rate limiting)",
)
bullet(doc, "Streaming responses for real-time conversation feel")
bullet(doc, "System prompt trained specifically for Vedic astrology scope")
bullet(
    doc,
    "Scope enforcement: politely declines off-topic queries and redirects to astrology",
)
bullet(
    doc,
    "Smart escalation: automatically suggests speaking with Guruji when personalized birth-chart analysis is needed",
)
bullet(doc, "Rate limiting: max 20 messages per minute per visitor")
bullet(doc, "Server-side proxy to protect the Gemini API key")
bullet(doc, "After free tier: gentle upsell to paid consultation via phone/WhatsApp")

heading2(doc, "4.7 Instant Palm Reading (/palm-reading)")
body(
    doc,
    (
        "An innovative feature allowing users to upload or capture a photo of their palm for instant "
        "AI-powered palmistry analysis."
    ),
)
bullet(doc, "Image upload / camera capture support")
bullet(doc, "AI analysis via Google Gemini Vision API")
bullet(
    doc,
    "Palmistry-specific analysis covering: life line, heart line, head line, fate line, mounts",
)
bullet(doc, "Personalized insights presented in a user-friendly format")
bullet(doc, "Escalation to Guruji for deeper analysis")

heading2(doc, "4.8 Testimonials (/testimonials)")
body(doc, "A dedicated page showcasing client success stories and reviews.")
bullet(doc, "9 curated testimonials from real-sounding clients across Bengaluru")
bullet(
    doc,
    "Filter by service type: Love, Marriage, Career, Finance, Homam, Birth Chart, Chat",
)
bullet(doc, "5-star ratings and client avatar initials")
bullet(doc, "Date stamps for authenticity")
bullet(doc, "Featured testimonials highlighted")
bullet(doc, "Smooth filtering animations")

heading2(doc, "4.9 Contact Us (/contact-us)")
bullet(
    doc,
    "Comprehensive enquiry form with fields: name, phone, email, birth details, service selection, message",
)
bullet(doc, "Form validation using React Hook Form + Zod schema")
bullet(doc, "Submission to backend API endpoint (/api/enquiry)")
bullet(doc, "WhatsApp direct link for instant messaging")
bullet(doc, "Google Maps embed for physical location")
bullet(doc, "Contact details: phone, email, address, working hours")
bullet(doc, "Social media links (Facebook, Instagram)")
bullet(doc, "Floating WhatsApp button (persistent across all pages)")

heading2(doc, "4.10 Legal Pages")
body(doc, "Four legally compliant pages ensuring transparency and trust:")
bullet(
    doc,
    "Disclaimer (/disclaimer) — Astrology guidance disclaimer, professional advice notice, personal responsibility",
)
bullet(
    doc,
    "Terms & Conditions (/terms-and-conditions) — Acceptance of terms, service descriptions, no guaranteed results, booking & payments",
)
bullet(
    doc,
    "Privacy Policy (/privacy-policy) — Data collection, usage, confidentiality, payment handling, data retention, user rights",
)
bullet(
    doc,
    "Refund & Cancellation (/refund-cancellation) — Cancellation windows (24h), refund eligibility, refund process, no-show policy, payment disputes",
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — ADMIN CMS
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "5. Admin CMS (Content Management System)")
add_divider(doc)

body(
    doc,
    (
        "A password-protected admin panel at /admin/* that allows the site owner to manage all "
        "content without touching code. Authentication is handled via Supabase Auth."
    ),
)

heading2(doc, "5.1 Admin Pages")
make_table(
    doc,
    ["Page", "Route", "Functionality"],
    [
        ["Admin Login", "/admin/login", "Secure login using Supabase Auth credentials"],
        [
            "Dashboard",
            "/admin/dashboard",
            "Overview cards for all manageable content sections",
        ],
        [
            "Services List",
            "/admin/services",
            "View, edit, and manage all astrology consultation services",
        ],
        [
            "New Service",
            "/admin/services/new",
            "Create a new astrology consultation service",
        ],
        [
            "Edit Service",
            "/admin/services/:slug",
            "Edit an existing service (title, description, pricing, FAQs, remedies)",
        ],
        ["Homams List", "/admin/homams", "View, edit, and manage all homam offerings"],
        ["New Homam", "/admin/homams/new", "Create a new homam ritual entry"],
        [
            "Edit Homam",
            "/admin/homams/:slug",
            "Edit an existing homam (details, pricing, booking info)",
        ],
        [
            "Page Content",
            "/admin/pages/:page",
            "Edit hero copy, content & FAQs for birth chart, chat, and palm reading pages",
        ],
    ],
    col_widths=[2.0, 2.5, 3.0],
)

heading2(doc, "5.2 Admin Features")
bullet(doc, "Supabase-powered authentication — secure login/logout")
bullet(doc, "Full CRUD (Create, Read, Update, Delete) for services and homams")
bullet(doc, "Content editing for dynamic pages (birth chart, chat, palm reading)")
bullet(
    doc,
    "Dashboard with summary cards showing count of services, homams, and page sections",
)
bullet(doc, "Protected routes — unauthenticated users are redirected to login")
bullet(doc, "Admin layout with consistent sidebar navigation")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — BACKEND API SERVICES
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "6. Backend API Services")
add_divider(doc)

body(doc, "The Express.js server (port 3001) exposes the following API endpoints:")

heading2(doc, "6.1 Public API Routes")
make_table(
    doc,
    ["Endpoint", "Method", "Purpose"],
    [
        [
            "/api/chat",
            "POST",
            "AI chatbot endpoint — proxies messages to Google Gemini with streaming response",
        ],
        [
            "/api/palm-reading",
            "POST",
            "Palm reading AI — sends palm image to Gemini Vision API",
        ],
        [
            "/api/enquiry",
            "POST",
            "Contact form submission — captures enquiry data and forwards to Guruji",
        ],
    ],
    col_widths=[2.0, 1.0, 3.5],
)

heading2(doc, "6.2 Admin API Routes")
make_table(
    doc,
    ["Endpoint", "Method", "Purpose"],
    [
        [
            "/api/admin/seed",
            "POST",
            "Initialize Supabase tables with default data (services, homams, pages)",
        ],
        [
            "/api/admin/services",
            "GET/POST",
            "Fetch all services or create a new service",
        ],
        [
            "/api/admin/services/:slug",
            "GET/PUT/DELETE",
            "Fetch, update, or delete a specific service",
        ],
        ["/api/admin/homams", "GET/POST", "Fetch all homams or create a new homam"],
        [
            "/api/admin/homams/:slug",
            "GET/PUT/DELETE",
            "Fetch, update, or delete a specific homam",
        ],
        [
            "/api/admin/pages/:page",
            "GET/PUT",
            "Fetch or update page content for dynamic pages",
        ],
    ],
    col_widths=[2.0, 1.0, 3.5],
)

heading2(doc, "6.3 Server-Side Features")
bullet(doc, "Rate limiting on AI endpoints (20 messages/min per visitor)")
bullet(doc, "Visitor tracking via HTTP-only cookies for free-tier AI chat")
bullet(doc, "Free-tier quota management (3 free AI questions per visitor)")
bullet(doc, "Server-side streaming of Gemini responses to the client")
bullet(doc, "CORS configuration for cross-origin requests")
bullet(doc, "Hot-reload during development using TSX watch mode")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — DATABASE & AUTHENTICATION
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "7. Database & Authentication (Supabase)")
add_divider(doc)

body(
    doc,
    (
        "Supabase provides the PostgreSQL database and authentication infrastructure. "
        "The application uses the Supabase JS client for both server-side (SSR) and client-side operations."
    ),
)

heading2(doc, "7.1 Database Tables")
make_table(
    doc,
    ["Table", "Purpose"],
    [
        [
            "services",
            "Stores all astrology consultation service entries (title, description, pricing, analysis, remedies, FAQs, ordering)",
        ],
        [
            "homams",
            "Stores all homam ritual entries (name, description, benefits, duration, pooja items, FAQs, booking instructions)",
        ],
        [
            "pages",
            "Stores editable content for dynamic pages (birth chart, chat, palm reading hero sections and FAQs)",
        ],
        [
            "testimonials",
            "Stores client testimonial entries (name, rating, service type, text, date, avatar)",
        ],
        ["users", "Supabase Auth managed — stores admin user accounts for the CMS"],
    ],
    col_widths=[2.0, 4.5],
)

heading2(doc, "7.2 Authentication")
bullet(doc, "Supabase Auth for admin login (email/password)")
bullet(doc, "Protected admin routes with authentication middleware")
bullet(doc, "Row Level Security (RLS) policies for data protection")
bullet(doc, "Session management with JWT tokens")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — KEY INTEGRATIONS
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "8. Key Integrations")
add_divider(doc)

heading2(doc, "8.1 Google Gemini AI")
body(doc, "The Gemini API is used for two major AI features:")
bullet(
    doc,
    "AI Chatbot — Generates context-aware responses for the Chat with Guruji feature",
)
bullet(
    doc,
    "Palm Reading — Analyzes palm images using Gemini's multimodal (vision) capabilities",
)
body(doc, "Key implementation details:")
bullet(doc, "System prompts enforce strict astrological scope")
bullet(doc, "Streaming responses for real-time conversation experience")
bullet(doc, "Free-tier visitor tracking with cookie-based identification")
bullet(doc, "Rate limiting to prevent API abuse")
bullet(doc, "Server-side proxy to protect API key from exposure")

heading2(doc, "8.2 WhatsApp Integration")
body(doc, "WhatsApp is deeply integrated as the primary communication channel:")
bullet(doc, "Direct WhatsApp links (wa.me) for instant messaging")
bullet(doc, "Floating WhatsApp button on all public pages")
bullet(doc, "WhatsApp as the primary contact method for bookings and enquiries")
bullet(doc, "Pre-formatted messages for quick contact initiation")

heading2(doc, "8.3 Payment Integration")
bullet(doc, "Razorpay integration for online payments")
bullet(doc, "UPI payment support")
bullet(doc, "Manual payment via screenshot upload for offline flexibility")
bullet(doc, "Payment references tracked in enquiry submissions")

heading2(doc, "8.4 Supabase")
bullet(doc, "PostgreSQL database for content storage")
bullet(doc, "Supabase Auth for admin authentication")
bullet(doc, "Supabase SSR client for server-side data fetching")
bullet(doc, "Supabase JS client for client-side queries")
bullet(doc, "Edge Functions support (configurable)")
bullet(doc, "Database migrations support")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — SEO & PERFORMANCE FEATURES
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "9. SEO & Performance Features")
add_divider(doc)

bullet(doc, "React Helmet Async for dynamic meta tag management per page")
bullet(doc, "Custom page titles and descriptions for every route")
bullet(doc, "Open Graph (OG) meta tags for social media sharing")
bullet(doc, "Google Fonts preconnect optimization (Cinzel, Cormorant Garamond, Inter)")
bullet(doc, "Semantic HTML structure")
bullet(doc, "Sitemap generation script (/scripts/generate-sitemap.ts)")
bullet(doc, "PWA support via service worker (sw.js)")
bullet(doc, "Lazy loading of components and images where applicable")
bullet(doc, "Font loading optimization with display=swap")
bullet(doc, "Vite's tree-shaking and code splitting for optimized bundles")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — DESIGN & UX FEATURES
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "10. Design & UX Features")
add_divider(doc)

heading2(doc, "10.1 Design System")
bullet(
    doc,
    "Tailwind CSS v4 with custom design tokens (gold/earth-toned palette matching the spiritual brand)",
)
bullet(
    doc,
    "Three Google Fonts: Cinzel (headings), Cormorant Garamond (serif body), Inter (UI text)",
)
bullet(
    doc,
    "Custom color scheme: earth tones, gold accents (#b67a1b), warm cream backgrounds (#faf4e8)",
)
bullet(doc, "Consistent border radius (rounded-3xl cards), shadows, and hover effects")

heading2(doc, "10.2 Animations & Transitions")
bullet(doc, "Page transition animations using Framer Motion")
bullet(doc, "Scroll-triggered reveal animations on the homepage")
bullet(doc, "Hover effects on cards and interactive elements")
bullet(doc, "Smooth filtering animations for testimonial section")
bullet(doc, "Floating animation for the WhatsApp button")

heading2(doc, "10.3 Responsive Design")
bullet(doc, "Mobile-first responsive layout using Tailwind breakpoints")
bullet(doc, "Collapsible mobile navigation with hamburger menu")
bullet(doc, "Responsive grids that adapt from 1 column (mobile) to 3 columns (desktop)")
bullet(doc, "Touch-friendly button sizes for mobile interaction")
bullet(doc, "Tested across desktop, tablet, and mobile viewports")

heading2(doc, "10.4 Accessibility")
bullet(doc, "Semantic HTML structure")
bullet(doc, "Proper heading hierarchy (h1 → h2 → h3)")
bullet(doc, "Color contrast compliance for readability")
bullet(doc, "Keyboard navigation support")
bullet(doc, "Alt text for meaningful images")
bullet(doc, "Form labels and validation messages")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — DEPLOYMENT & BUILD
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "11. Deployment & Build Configuration")
add_divider(doc)

heading2(doc, "11.1 Build Scripts")
make_table(
    doc,
    ["Script", "Command", "Purpose"],
    [
        [
            "dev",
            "npm run dev",
            "Start development server (Vite + Express concurrently)",
        ],
        ["dev:vite", "npm run dev:vite", "Start only Vite dev server"],
        ["dev:api", "npm run dev:api", "Start only Express API server"],
        ["build", "npm run build", "TypeScript compile + Vite production build"],
        ["build:all", "npm run build:all", "Build frontend + server bundle"],
        ["start", "npm run start", "Start Express server in production mode"],
        [
            "server",
            "npm run server",
            "Start Express server with hot reload (TSX watch)",
        ],
        ["preview", "npm run preview", "Preview production build locally"],
        ["lint", "npm run lint", "Run ESLint on source files"],
        ["typecheck", "npm run typecheck", "Run TypeScript type checking"],
        [
            "generate-sitemap",
            "npm run generate-sitemap",
            "Generate XML sitemap for SEO",
        ],
    ],
    col_widths=[1.5, 1.8, 3.2],
)

heading2(doc, "11.2 Build Configuration")
bullet(doc, "TypeScript strict mode enabled")
bullet(doc, "Path aliases configured: @/*, $lib/*, $data/*")
bullet(doc, "Vite proxy: /api routes forwarded to localhost:3001 during development")
bullet(doc, "Production output to /dist directory")
bullet(doc, "Express serves static files from /dist in production")

heading2(doc, "11.3 Engine Requirements")
make_table(
    doc,
    ["Requirement", "Version"],
    [
        ["Node.js", ">= 22"],
        ["npm", ">= 10"],
    ],
    col_widths=[3.0, 3.5],
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12 — ENVIRONMENT VARIABLES
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "12. Environment Variables & Configuration")
add_divider(doc)

make_table(
    doc,
    ["Variable", "Required", "Purpose"],
    [
        ["VITE_SUPABASE_URL", "Yes", "Supabase project URL"],
        ["VITE_SUPABASE_ANON_KEY", "Yes", "Supabase anonymous/public API key"],
        ["VITE_SUPABASE_FUNCTIONS_URL", "No", "Override Supabase edge functions URL"],
        ["GEMINI_API_KEY", "Yes (server)", "Google Gemini API key for AI features"],
        [
            "VITE_API_MODE",
            "No",
            "API routing mode: 'local' (Express) or 'supabase' (Edge Functions)",
        ],
        [
            "VITE_SITE_URL",
            "No",
            "Production site URL for canonical links and redirects",
        ],
        ["VITE_WHATSAPP_NUMBER", "No", "WhatsApp number (defaults to +91 98861 00565)"],
    ],
    col_widths=[2.5, 1.5, 3.5],
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 13 — TESTING & QUALITY ASSURANCE
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "13. Testing & Quality Assurance")
add_divider(doc)

heading2(doc, "13.1 Code Quality")
bullet(doc, "TypeScript strict mode — full type safety across frontend and server")
bullet(doc, "ESLint for code style enforcement and error detection")
bullet(doc, "Zod runtime validation on all form inputs and API request bodies")
bullet(doc, "React Hook Form for controlled, validated form state")

heading2(doc, "13.2 Error Handling")
bullet(doc, "Graceful degradation when Gemini API is unavailable")
bullet(doc, "Error boundaries for React component failures")
bullet(doc, "API error responses with user-friendly messages")
bullet(doc, "Form validation with clear inline error messages")
bullet(doc, "Rate limiting with friendly retry messages")

heading2(doc, "13.3 Manual Testing Coverage")
bullet(doc, "All 14 public pages verified for rendering and responsiveness")
bullet(doc, "Admin CMS login and CRUD operations tested")
bullet(
    doc, "AI chat flow tested: free tier, rate limiting, scope enforcement, escalation"
)
bullet(doc, "Palm reading image upload and analysis tested")
bullet(doc, "Contact form submission and validation tested")
bullet(doc, "WhatsApp links verified for all platforms")
bullet(doc, "Legal pages verified for completeness")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 14 — SUMMARY OF DELIVERABLES
# ══════════════════════════════════════════════════════════════════════════════
heading1(doc, "14. Summary of Deliverables")
add_divider(doc)

body(
    doc,
    (
        "The following is a consolidated list of everything that has been built and delivered "
        "as part of this project:"
    ),
)

heading2(doc, "14.1 Public Website (14 Pages)")
make_table(
    doc,
    ["#", "Page", "Route", "Key Feature"],
    [
        [
            "1",
            "Homepage",
            "/",
            "Hero, services overview, testimonials, FAQ, AI chat widget",
        ],
        ["2", "About Us", "/about-us", "Biography, credentials, philosophy, location"],
        [
            "3",
            "Services",
            "/services",
            "10 consultation service listings with detailed info",
        ],
        [
            "4",
            "Service Detail",
            "/services/:slug",
            "Full service page with enquiry form",
        ],
        ["5", "Homams", "/homams", "18 homam listings with booking info"],
        ["6", "Homam Detail", "/homams/:slug", "Full homam page with booking form"],
        ["7", "Birth Chart PDF", "/birth-chart-pdf", "Kundli PDF request form"],
        [
            "8",
            "Chat with Guruji",
            "/chat-with-guruji",
            "AI chatbot with free-tier quota",
        ],
        ["9", "Palm Reading", "/palm-reading", "AI-powered palm image analysis"],
        ["10", "Testimonials", "/testimonials", "Client reviews with filtering"],
        [
            "11",
            "Contact Us",
            "/contact-us",
            "Enquiry form, WhatsApp, map, social links",
        ],
        ["12", "Disclaimer", "/disclaimer", "Legal disclaimer page"],
        ["13", "Terms & Conditions", "/terms-and-conditions", "Legal terms page"],
        ["14", "Privacy Policy", "/privacy-policy", "Data privacy page"],
    ],
    col_widths=[0.5, 1.5, 1.8, 3.7],
)

heading2(doc, "14.2 Admin CMS (8 Pages)")
make_table(
    doc,
    ["#", "Page", "Route", "Key Feature"],
    [
        ["1", "Admin Login", "/admin/login", "Supabase Auth login"],
        ["2", "Dashboard", "/admin/dashboard", "Overview of all manageable content"],
        ["3", "Services List", "/admin/services", "CRUD for consultation services"],
        ["4", "New Service", "/admin/services/new", "Create new service"],
        ["5", "Edit Service", "/admin/services/:slug", "Edit existing service"],
        ["6", "Homams List", "/admin/homams", "CRUD for homam entries"],
        ["7", "New Homam", "/admin/homams/new", "Create new homam"],
        ["8", "Edit Homam", "/admin/homams/:slug", "Edit existing homam"],
    ],
    col_widths=[0.5, 1.5, 2.0, 3.5],
)

heading2(doc, "14.3 AI-Powered Features")
bullet(
    doc,
    "AI Chatbot (Google Gemini) — free-tier chat with 3 free messages, streaming responses, astrological scope enforcement, and smart escalation to Guruji",
)
bullet(
    doc,
    "Palm Reading (Google Gemini Vision) — upload palm photo for instant AI palmistry analysis",
)

heading2(doc, "14.4 Reusable UI Components")
bullet(doc, "Navbar — responsive with mobile hamburger menu")
bullet(doc, "Footer — sitemap links, social links, legal links")
bullet(doc, "FloatingWhatsApp — persistent CTA button")
bullet(doc, "PageTransition — animated page transitions via Framer Motion")
bullet(doc, "ServiceAiChat — embedded AI chat widget")
bullet(doc, "AiChatWidget — standalone chat component")
bullet(doc, "PalmReader — palm image upload and analysis component")
bullet(doc, "AskGurujiButton — quick AI chat trigger")
bullet(doc, "Cards — service cards, homam cards, testimonial cards")
bullet(doc, "Forms — enquiry form, contact form components")
bullet(doc, "FAQ accordion — collapsible FAQ sections")
bullet(doc, "SEO component — React Helmet wrapper")

heading2(doc, "14.5 Backend API Endpoints")
bullet(doc, "POST /api/chat — AI chatbot with streaming")
bullet(doc, "POST /api/palm-reading — Palm image analysis")
bullet(doc, "POST /api/enquiry — Contact form submission")
bullet(doc, "GET/POST /api/admin/services — Service CRUD")
bullet(doc, "GET/PUT/DELETE /api/admin/services/:slug — Single service CRUD")
bullet(doc, "GET/POST /api/admin/homams — Homam CRUD")
bullet(doc, "GET/PUT/DELETE /api/admin/homams/:slug — Single homam CRUD")
bullet(doc, "GET/PUT /api/admin/pages/:page — Page content management")

heading2(doc, "14.6 Database Schema")
bullet(doc, "services table — 10 astrology consultation entries")
bullet(doc, "homams table — 18 sacred ritual entries")
bullet(doc, "pages table — Dynamic page content (birth chart, chat, palm reading)")
bullet(doc, "testimonials table — 9 client review entries")
bullet(doc, "users table — Admin CMS user accounts (Supabase Auth)")

heading2(doc, "14.7 Static Data Provided")
bullet(
    doc,
    "10 Astrology Services — each with title, description, problem statement, analysis points, remedies, benefits, pricing, FAQs",
)
bullet(
    doc,
    "18 Homam Rituals — each with name, description, benefits, suitable-for list, pooja items, booking instructions, FAQs",
)
bullet(doc, "9 Client Testimonials — with ratings, service types, dates, and locations")
bullet(
    doc,
    "FAQs for 5 sections — homepage, birth chart, chat, contact, and per-service FAQs",
)
bullet(
    doc,
    "4 Legal Pages — Privacy Policy, Terms & Conditions, Disclaimer, Refund & Cancellation Policy",
)

heading2(doc, "14.8 SEO & PWA")
bullet(doc, "Dynamic meta tags via React Helmet Async")
bullet(doc, "Sitemap generation script")
bullet(doc, "PWA service worker for offline capability")
bullet(doc, "Open Graph meta tags for social sharing")
bullet(doc, "Google Fonts preconnect optimization")

heading2(doc, "14.9 Development Tooling")
bullet(
    doc, "Custom dev orchestrator (scripts/dev.mjs) — runs Vite + Express concurrently"
)
bullet(doc, "Hot reload via TSX watch for server code")
bullet(doc, "Vite fast refresh for frontend components")
bullet(doc, "TypeScript strict mode with path aliases")
bullet(doc, "ESLint configuration")
bullet(
    doc, "Production build pipeline (TypeScript compile → Vite build → Express serve)"
)

# ── final divider ─────────────────────────────────────────────────────────────
add_divider(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
run = p.add_run("— End of Document —")
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Vedic Astrology — Sampath Kumara | Ancient Wisdom • Cosmic Guidance")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# Save
output_path = os.path.join(
    os.path.dirname(__file__), "vedic-astrology-project-documentation.docx"
)
doc.save(output_path)
print(f"Document saved to: {output_path}")
